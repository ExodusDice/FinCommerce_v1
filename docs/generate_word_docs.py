import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(99, 102, 241) # Indigo #6366f1
        p.paragraph_format.space_before = Pt(18)
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(6, 182, 212) # Cyan #06b6d4
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(55, 65, 81) # Dark Gray
        
    return p

def convert_md_to_docx(md_path, docx_path):
    print(f"Converting {md_path} to {docx_path}...")
    doc = Document()
    
    # Configure margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Read MD file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    for line in lines:
        stripped = line.strip()
        
        # Check if we exit table
        if in_table and (not stripped.startswith('|') or stripped.startswith('|---') or stripped == ''):
            if table_headers and table_rows:
                # Create table in Word
                cols = len(table_headers)
                t = doc.add_table(rows=1, cols=cols)
                t.autofit = True
                
                # Header row
                hdr_cells = t.rows[0].cells
                for i, text in enumerate(table_headers):
                    hdr_cells[i].text = text.strip()
                    set_cell_background(hdr_cells[i], '6366F1') # Indigo Header
                    set_cell_margins(hdr_cells[i])
                    run = hdr_cells[i].paragraphs[0].runs[0]
                    run.font.name = 'Segoe UI'
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.size = Pt(9.5)
                
                # Data rows
                for r_idx, r_data in enumerate(table_rows):
                    row_cells = t.add_row().cells
                    bg_color = 'F9FAFB' if r_idx % 2 == 1 else 'FFFFFF' # Zebra striping
                    for c_idx, text in enumerate(r_data):
                        if c_idx < len(row_cells):
                            row_cells[c_idx].text = text.strip()
                            set_cell_background(row_cells[c_idx], bg_color)
                            set_cell_margins(row_cells[c_idx])
                            if row_cells[c_idx].paragraphs:
                                run = row_cells[c_idx].paragraphs[0].runs[0] if row_cells[c_idx].paragraphs[0].runs else row_cells[c_idx].paragraphs[0].add_run()
                                run.font.name = 'Segoe UI'
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(55, 65, 81)
                
                doc.add_paragraph() # Add empty spacing after table
                
            in_table = False
            table_headers = []
            table_rows = []
            
        # Parse MD formats
        if stripped.startswith('# '):
            add_styled_heading(doc, stripped[2:], 1)
        elif stripped.startswith('## '):
            add_styled_heading(doc, stripped[3:], 2)
        elif stripped.startswith('### '):
            add_styled_heading(doc, stripped[4:], 3)
        elif stripped.startswith('|'):
            # It's a table row
            if stripped.startswith('|---') or stripped.startswith('| :---') or stripped.startswith('|:---'):
                # Divider line, ignore
                continue
            cells = [c.strip() for c in stripped.split('|')[1:-1]]
            if not in_table:
                in_table = True
                table_headers = cells
            else:
                table_rows.append(cells)
        elif stripped.startswith('* ') or stripped.startswith('- '):
            # Bullet list
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(stripped[2:])
            run.font.name = 'Segoe UI'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(55, 65, 81)
        elif stripped.startswith('1. ') or stripped.startswith('2. ') or re.match(r'^\d+\.\s', stripped):
            # Numbered list
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            content = re.sub(r'^\d+\.\s', '', stripped)
            run = p.add_run(content)
            run.font.name = 'Segoe UI'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(55, 65, 81)
        elif stripped.startswith('> '):
            # Blockquote
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[2:])
            run.font.name = 'Segoe UI'
            run.font.italic = True
            run.font.color.rgb = RGBColor(99, 102, 241)
            run.font.size = Pt(10)
        elif stripped == '':
            # Empty line
            continue
        else:
            # Paragraph
            # Strip markdown link syntax but keep label text (e.g. [text](file://...) -> text)
            clean_text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(clean_text)
            run.font.name = 'Segoe UI'
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(55, 65, 81)
            
    doc.save(docx_path)
    print(f"Generated {docx_path} successfully!")

if __name__ == "__main__":
    docs_dir = os.path.dirname(__file__)
    files_to_convert = [
        ("BRD.md", "BRD.docx"),
        ("FSD.md", "FSD.docx"),
        ("Technical_Spec.md", "Technical_Spec.docx"),
        ("Test_Cases.md", "Test_Cases.docx"),
        ("DFD.md", "DFD.docx"),
        ("Flowchart.md", "Flowchart.docx"),
        ("ERD.md", "ERD.docx"),
        ("UC_Diagram.md", "UC_Diagram.docx"),
        ("Sequential_Diagram.md", "Sequential_Diagram.docx"),
        ("API_Specs.md", "API_Specs.docx"),
        ("Performance_Test_Plan.md", "Performance_Test_Plan.docx"),
        ("Compatibility_Test_Plan.md", "Compatibility_Test_Plan.docx"),
        ("Project_Timeline.md", "Project_Timeline.docx"),
        ("User_Manual.md", "User_Manual.docx"),
        ("CHANGELOG.md", "CHANGELOG.docx"),
        ("Subscription_Test_Plan.md", "Subscription_Test_Plan.docx")
    ]
    for md_file, docx_file in files_to_convert:
        md_p = os.path.join(docs_dir, md_file)
        docx_p = os.path.join(docs_dir, docx_file)
        if os.path.exists(md_p):
            convert_md_to_docx(md_p, docx_p)
        else:
            print(f"Warning: {md_file} not found.")
